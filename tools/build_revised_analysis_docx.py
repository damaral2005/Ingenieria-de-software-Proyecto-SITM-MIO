from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"docs\Analisis_de_experimento-PARCIALFINAL-version3-revisado.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2EC")


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_table_borders(t)
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "E8EEF5")
        if widths:
            cell.width = Inches(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return t


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181 if level < 3 else 120)
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        rest = text[len(bold_prefix):]
        p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_code(doc, code):
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        run.font.size = Pt(8)
    doc.add_paragraph()


def fmt_s(ms):
    return "{:.3f} s".format(ms / 1000.0)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Universidad Icesi - Ingenieria de Software IV")
    r.bold = True
    r.font.size = Pt(14)

    for text, size, bold in [
        ("Proyecto Final - SITM-MIO", 18, True),
        ("Documento de Resultados del Experimento", 16, True),
        ("Calculo de velocidades promedio por ruta y mes", 13, False),
        ("Comparativa V1 monolitica, V2 concurrente y V3 distribuida con ZeroC Ice", 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold

    doc.add_paragraph()
    add_para(doc, "Sistema bajo estudio: SITM-MIO")
    add_para(doc, "Despliegue del experimento: datos reales MiniPilot y Pilot, incluyendo ejecuciones V3 local/file-based y V3 con ZeroC Ice en multiples PCs.")
    add_para(doc, "Integrantes: Juan Jose Reyes Ramos, Anderson Romero, Daniel, Dayanna.")
    add_para(doc, "Revision preparada con evidencia de las ramas main, version-2-threadpool y codex/version-3-distributed-master-worker.")
    doc.add_page_break()

    add_heading(doc, "Resumen ejecutivo", 1)
    add_para(
        doc,
        "Este documento actualiza el analisis experimental del calculo de velocidades promedio por ruta y mes del SITM-MIO usando la evidencia real disponible para las tres versiones. La comparacion se enfoca en dos datasets: MiniPilot, usado para comparar rendimiento y correctitud sobre una muestra manejable, y Pilot, el dataset completo de aproximadamente 67 GB y 806,400,773 datagramas.",
    )
    add_para(
        doc,
        "La conclusion principal cambia de una lectura puramente de rendimiento a una lectura de viabilidad arquitectonica. V3 local fue la ejecucion mas rapida en MiniPilot, pero V3 con Ice fue mas lenta por el costo de serializacion, red y decodificacion de particiones. Aun asi, Ice es la evidencia mas fuerte de distribucion real porque el master coordino workers remotos.",
    )
    add_bullets(
        doc,
        [
            "Correctitud estructural: las ejecuciones MiniPilot de V1, V2 y V3 producen 111 filas de salida y no reportan valores NaN, Infinity o null. No se debe afirmar igualdad bit a bit si no se adjunta un diff formal.",
            "Limite de memoria: V1 y V2 fallan con OutOfMemoryError en Pilot; la concurrencia local no elimina el problema de cargar demasiados datos en un solo JVM.",
            "Escalabilidad: V3 completa Pilot tanto en modo distribuido local/file-based como con ZeroC Ice, produciendo 1,443 filas y 736,951,733 segmentos validos.",
            "Costo de distribucion: en V3 el particionamiento domina el tiempo total. En Ice, el sobrecosto de red hace que el tiempo aumente frente al modo local.",
        ],
    )

    add_heading(doc, "Que se cambio frente al analisis anterior", 1)
    table(
        doc,
        ["Tipo", "Cambio aplicado", "Razon"],
        [
            ["Se corrigio", "La metrica base de V1 MiniPilot se dejo en 64,081 ms.", "Es la evidencia limpia que usa la misma limpieza de datos que V2/V3."],
            ["Se corrigio", "V3 no se presenta como siempre mas rapida: se separa V3 local de V3 Ice.", "Ice completo MiniPilot tarda 53,780 ms y Pilot 84.99 min por costo de red/serializacion."],
            ["Se agrego", "Resultados de Pilot completo con V3 local y V3 Ice.", "Son los datos que validan escalabilidad real sobre 806,400,773 datagramas."],
            ["Se agrego", "Comparacion de V2 por threads y evidencia de fallo en 70 GB.", "La rama version-2-threadpool registra 39,201 ms con 4 threads, 38,643 ms con 12 threads y OOM en Pilot."],
            ["Se matizo", "La correctitud se expresa como equivalencia estructural y conteos compatibles.", "No hay en el repo una comparacion bit a bit formal registrada entre todas las salidas."],
            ["Se reemplazo", "Las graficas embebidas se reemplazaron por scripts Python reproducibles.", "Permite regenerar las figuras con los datos reales de las ramas."],
        ],
        widths=[1.1, 2.4, 3.0],
    )

    add_heading(doc, "Introduccion", 1)
    add_para(
        doc,
        "El experimento evalua tres alternativas arquitectonicas para procesar datagramas GPS historicos del SITM-MIO y calcular velocidades promedio por ruta y mes. Todas las versiones mantienen la misma salida CSV: route_id, month, total_distance_km, total_time_hours, avg_speed_kmh, avg_segment_speed_kmh, valid_segments y buses_observed.",
    )
    table(
        doc,
        ["Version", "Arquitectura", "Proposito"],
        [
            ["V1", "Monolitica secuencial", "Linea base funcional y evidencia del limite de memoria."],
            ["V2", "Concurrente con Thread Pool", "Reducir tiempo con paralelismo local en un solo JVM."],
            ["V3 local", "Distributed Master-Worker file-based", "Particionar datos y procesar work units acotados en memoria."],
            ["V3 Ice", "Distributed Master-Worker con ZeroC Ice", "Probar despliegue distribuido real con workers remotos."],
        ],
        widths=[1.0, 2.0, 3.5],
    )
    add_para(
        doc,
        "El calculo no usa odometer porque no es confiable para este caso. La distancia se calcula con Haversine a partir de coordenadas GPS y la velocidad se deriva de la diferencia temporal entre puntos consecutivos del mismo bus y ruta.",
    )

    add_heading(doc, "Entorno experimental y fuentes", 1)
    table(
        doc,
        ["Fuente", "Dato usado"],
        [
            ["main / commit 865882f", "V1 MiniPilot limpio: 64,081 ms, 7,896,735 datagramas limpios, 7,494,051 segmentos validos."],
            ["origin/main / commit 082f88a", "V1 Pilot 70 GB falla por OutOfMemoryError; build fallido en 7m40s."],
            ["version-2-threadpool / commit d7a9070", "V2 MiniPilot por numero de threads: 1, 2, 4, 6, 8 y 12 threads."],
            ["origin/version-2-threadpool / commit 4f07da1", "V2 Pilot 70 GB falla por OutOfMemoryError; build fallido en 7m20s."],
            ["codex/version-3-distributed-master-worker", "V3 MiniPilot, V3 Pilot local/file-based, y ejecuciones ZeroC Ice MiniPilot/Pilot."],
        ],
        widths=[2.2, 4.3],
    )

    add_heading(doc, "Resultados MiniPilot", 1)
    add_para(
        doc,
        "MiniPilot contiene 8,145,462 datagramas reales. En la evidencia limpia, V1, V2 y V3 trabajan sobre 7,896,735 datagramas limpiados y reportan 7,494,051 segmentos validos. Esto permite comparar rendimiento sin mezclar cambios de limpieza.",
    )
    mini_rows = [
        ["V1 monolitica", "1 proceso / 1 JVM", "64,081", fmt_s(64081), "111", "7,494,051"],
        ["V2 ThreadPool", "4 threads", "39,201", fmt_s(39201), "111", "7,494,051"],
        ["V2 ThreadPool", "12 threads", "38,643", fmt_s(38643), "111", "7,494,051"],
        ["V3 local", "4 workers / 8 particiones", "30,814", fmt_s(30814), "111", "7,494,051"],
        ["V3 Ice", "2 workers remotos / 2 particiones", "53,780", fmt_s(53780), "111", "7,494,051"],
    ]
    table(
        doc,
        ["Version", "Configuracion", "Runtime ms", "Runtime", "Filas", "Segmentos validos"],
        mini_rows,
        widths=[1.25, 1.9, 1.0, 0.9, 0.65, 1.0],
    )
    add_para(
        doc,
        "En MiniPilot, V3 local es la configuracion mas rapida de las medidas disponibles: 30.814 s. V2 mejora mucho frente a V1, pero se estabiliza entre 4 y 12 threads, lo que sugiere saturacion por I/O y costos compartidos dentro del mismo JVM. V3 Ice no gana en tiempo sobre V3 local porque paga serializacion, transferencia de particiones por red y decodificacion en workers remotos; su valor principal es demostrar el despliegue distribuido real.",
    )

    add_heading(doc, "Grafica 1 - Comparacion de rendimiento MiniPilot", 2)
    add_para(doc, "Script Python para generar la grafica en este punto del documento:")
    add_code(
        doc,
        r"""
import matplotlib.pyplot as plt

labels = ["V1\nmonolitica", "V2\n4 threads", "V2\n12 threads", "V3 local\n4w/8p", "V3 Ice\n2w/2p"]
runtime_ms = [64081, 39201, 38643, 30814, 53780]
runtime_s = [v / 1000 for v in runtime_ms]

plt.figure(figsize=(8, 4.5))
bars = plt.bar(labels, runtime_s, color=["#7f8c8d", "#3498db", "#2980b9", "#27ae60", "#f39c12"])
plt.ylabel("Tiempo total (segundos)")
plt.title("MiniPilot: runtime total por arquitectura")
for bar, value in zip(bars, runtime_s):
    plt.text(bar.get_x() + bar.get_width() / 2, value + 1, f"{value:.1f}s", ha="center", fontsize=9)
plt.tight_layout()
plt.savefig("grafica_1_minipilot_runtime.png", dpi=180)
plt.show()
""",
    )

    add_heading(doc, "Grafica 2 - Throughput y nivel de concurrencia", 2)
    add_para(doc, "Script Python para generar la grafica en este punto del documento:")
    add_code(
        doc,
        r"""
import matplotlib.pyplot as plt

raw_datagrams = 8_145_462
threads = [1, 2, 4, 6, 8, 12]
v2_runtime_ms = [138309, 52587, 39201, 39447, 39086, 38643]
v2_throughput = [raw_datagrams / (ms / 1000) for ms in v2_runtime_ms]
v1_throughput = raw_datagrams / (64081 / 1000)
v3_local_throughput = raw_datagrams / (30814 / 1000)
v3_ice_throughput = raw_datagrams / (53780 / 1000)

plt.figure(figsize=(8, 4.5))
plt.plot(threads, v2_throughput, marker="o", label="V2 ThreadPool")
plt.axhline(v1_throughput, linestyle="--", color="#7f8c8d", label="V1 monolitica")
plt.axhline(v3_local_throughput, linestyle="--", color="#27ae60", label="V3 local")
plt.axhline(v3_ice_throughput, linestyle="--", color="#f39c12", label="V3 Ice")
plt.xlabel("Threads en V2")
plt.ylabel("Datagramas crudos por segundo")
plt.title("MiniPilot: throughput observado")
plt.legend()
plt.grid(axis="y", alpha=0.25)
plt.tight_layout()
plt.savefig("grafica_2_throughput.png", dpi=180)
plt.show()
""",
    )

    add_heading(doc, "Resultados Pilot completo", 1)
    add_para(
        doc,
        "Pilot es el dataset completo: 806,400,773 datagramas crudos y aproximadamente 67 GB. En este escenario aparece el limite real de arquitectura: V1 y V2 fallan por memoria, mientras que V3 completa el procesamiento al dividir el trabajo en particiones acotadas.",
    )
    table(
        doc,
        ["Version", "Estado en Pilot", "Configuracion / evidencia", "Resultado"],
        [
            ["V1 monolitica", "Falla", "Un JVM; carga demasiados datagramas limpiados en memoria.", "OutOfMemoryError; build fallido en 7m40s."],
            ["V2 ThreadPool", "Falla", "Un JVM con 8 threads; concurrencia local.", "OutOfMemoryError; build fallido en 7m20s."],
            ["V3 local", "Completa", "64 particiones file-based por routeId + busId.", "73.43 min; 1,443 filas; 736,951,733 segmentos validos."],
            ["V3 Ice", "Completa", "2 workers remotos; 1024 particiones enviadas por Ice.", "84.99 min; 1,443 filas; 736,951,733 segmentos validos."],
        ],
        widths=[1.1, 0.85, 2.5, 2.05],
    )
    add_para(
        doc,
        "La diferencia entre V3 local y V3 Ice no contradice la utilidad de Ice. El modo local procesa archivos de particion ya materializados; Ice debe enviar contenido de particion por RPC, serializarlo, transferirlo por red y reconstruirlo en cada worker. Por eso Ice es mas costoso en tiempo, pero prueba la topologia distribuida requerida.",
    )

    add_heading(doc, "Analisis interno de V3", 1)
    table(
        doc,
        ["Ejecucion V3", "Particiones", "Workers", "Particionamiento", "Workers", "Merge", "Total"],
        [
            ["MiniPilot local", "8", "4", "26,921 ms", "3,859 ms", "15 ms", "30,814 ms"],
            ["MiniPilot Ice", "2", "2 remotos", "30,570 ms", "23,134 ms", "47 ms", "53,780 ms"],
            ["Pilot local", "64", "particiones procesadas", "3,180,938 ms", "1,224,559 ms", "170 ms", "4,405,667 ms"],
            ["Pilot Ice", "1024", "2 remotos", "3,237,629 ms", "1,861,148 ms", "493 ms", "5,099,359 ms"],
        ],
        widths=[1.2, 0.75, 1.0, 1.1, 1.0, 0.75, 1.0],
    )
    add_para(
        doc,
        "El patron se repite en ambas escalas: el merge es despreciable porque los workers reducen millones de puntos GPS a agregados parciales compactos por ruta y mes. El costo dominante es leer, limpiar y particionar el CSV. En Ice, el tiempo de workers aumenta porque cada particion viaja como payload RPC y debe reconstruirse en el worker.",
    )

    add_heading(doc, "Grafica 3 - Desglose de tiempo V3", 2)
    add_para(doc, "Script Python para generar la grafica en este punto del documento:")
    add_code(
        doc,
        r"""
import matplotlib.pyplot as plt
import numpy as np

labels = ["Mini local", "Mini Ice", "Pilot local", "Pilot Ice"]
partition_ms = np.array([26921, 30570, 3180938, 3237629])
worker_ms = np.array([3859, 23134, 1224559, 1861148])
merge_ms = np.array([15, 47, 170, 493])

partition_min = partition_ms / 60000
worker_min = worker_ms / 60000
merge_min = merge_ms / 60000

x = np.arange(len(labels))
plt.figure(figsize=(9, 4.8))
plt.bar(x, partition_min, label="Particionamiento", color="#2c7fb8")
plt.bar(x, worker_min, bottom=partition_min, label="Workers", color="#41ab5d")
plt.bar(x, merge_min, bottom=partition_min + worker_min, label="Merge", color="#fdae61")
plt.xticks(x, labels)
plt.ylabel("Tiempo (minutos)")
plt.title("V3: costo por fase")
plt.legend()
plt.tight_layout()
plt.savefig("grafica_3_v3_fases.png", dpi=180)
plt.show()
""",
    )

    add_heading(doc, "Validacion de drivers de arquitectura", 1)
    table(
        doc,
        ["Driver", "Resultado experimental actualizado", "Estado"],
        [
            ["Performance", "V3 local reduce MiniPilot a 30.814 s. V2 mejora frente a V1, pero se estabiliza alrededor de 39 s. V3 Ice es mas lento que V3 local por costo de red.", "Cumple con matiz"],
            ["Scalability", "Solo V3 completa Pilot de 67 GB. V1 y V2 fallan por OutOfMemoryError.", "Cumple"],
            ["Correctness", "Todas las ejecuciones MiniPilot reportan 111 filas y 7,494,051 segmentos validos en la evidencia limpia. Falta registrar diff bit a bit.", "Cumple parcial"],
            ["Modifiability", "V3 conserva calculo, parsing y escritura separados de la infraestructura distribuida.", "Cumple"],
            ["Reliability / Availability", "El sistema detecta fallas de worker/proceso, pero no implementa reintentos automaticos ni recuperacion ante caida de worker.", "Parcial"],
        ],
        widths=[1.25, 4.4, 0.85],
    )

    add_heading(doc, "Punto de inflexion: cuando conviene distribuir", 1)
    add_para(
        doc,
        "La distribucion no debe justificarse solamente por ser mas rapida. En MiniPilot, V3 local gana en tiempo, pero V2 con ThreadPool ofrece una mejora importante con menor complejidad operacional. El punto fuerte de V3 aparece cuando el dataset supera lo que un JVM puede mantener en memoria de forma confiable.",
    )
    add_bullets(
        doc,
        [
            "Region monolitica: datasets pequenos o pruebas funcionales, donde la simplicidad de V1 puede ser suficiente.",
            "Region concurrente: datasets intermedios, donde V2 reduce tiempo sin incorporar red ni coordinacion remota.",
            "Region distribuida: datasets masivos como Pilot, donde V1/V2 fallan y V3 es la unica alternativa observada que completa el procesamiento.",
            "Region Ice: cuando el objetivo es demostrar o usar workers remotos reales; se acepta sobrecosto de comunicacion a cambio de despliegue distribuido autentico.",
        ],
    )

    add_heading(doc, "Grafica 4 - Viabilidad por tamano de dataset", 2)
    add_para(doc, "Script Python para generar la grafica en este punto del documento:")
    add_code(
        doc,
        r"""
import matplotlib.pyplot as plt

mini_gb = 0.687
pilot_gb = 67

series = {
    "V1 monolitica": {"x": [mini_gb], "y": [64081 / 60000], "color": "#7f8c8d"},
    "V2 4 threads": {"x": [mini_gb], "y": [39201 / 60000], "color": "#2980b9"},
    "V3 local": {"x": [mini_gb, pilot_gb], "y": [30814 / 60000, 4405667 / 60000], "color": "#27ae60"},
    "V3 Ice": {"x": [mini_gb, pilot_gb], "y": [53780 / 60000, 5099359 / 60000], "color": "#f39c12"},
}

plt.figure(figsize=(8, 4.8))
for label, data in series.items():
    plt.plot(data["x"], data["y"], marker="o", label=label, color=data["color"])

plt.scatter([pilot_gb, pilot_gb], [7.67, 7.33], marker="x", s=90, color=["#7f8c8d", "#2980b9"])
plt.text(pilot_gb, 8.6, "V1/V2: OOM", ha="right", fontsize=9)

plt.xscale("log")
plt.xlabel("Tamano aproximado del dataset (GB, escala log)")
plt.ylabel("Tiempo observado (minutos)")
plt.title("Viabilidad experimental por tamano de dataset")
plt.legend()
plt.grid(True, which="both", axis="both", alpha=0.25)
plt.tight_layout()
plt.savefig("grafica_4_viabilidad_dataset.png", dpi=180)
plt.show()
""",
    )

    add_heading(doc, "Amenazas a la validez y limitaciones", 1)
    add_bullets(
        doc,
        [
            "Las mediciones provienen de hosts y momentos de ejecucion diferentes; por tanto, los tiempos son evidencia experimental, no benchmark controlado de laboratorio.",
            "La version Ice usa 1024 particiones para Pilot porque 64 y 512 generaron presion de memoria o perdida de conexion; esto aumenta overhead frente al modo file-based.",
            "No se adjunta en la evidencia actual un diff formal de salidas V1/V2/V3 para afirmar identidad bit a bit.",
            "No se midio tolerancia a fallos con caida intencional de workers ni reintentos automaticos.",
            "El particionamiento sigue siendo fase dominante y mayormente secuencial; futuras versiones podrian estudiar particionamiento paralelo o streaming distribuido mas fino.",
        ],
    )

    add_heading(doc, "Conclusiones", 1)
    add_para(
        doc,
        "La evidencia actualizada muestra que V3 es la arquitectura necesaria para el caso operativo de mayor volumen del SITM-MIO. V1 y V2 son utiles para validar funcionalidad y comparar rendimiento en MiniPilot, pero ambas fallan con OutOfMemoryError en Pilot. V3, al particionar por routeId + busId, mantiene juntos los puntos necesarios para calcular segmentos correctos y reduce el problema a unidades de trabajo acotadas.",
    )
    add_para(
        doc,
        "La comparacion entre V3 local y V3 Ice debe leerse con cuidado: V3 local es mas eficiente en tiempo, mientras que V3 Ice es mas representativa del despliegue distribuido real exigido por la version 3. El costo extra de Ice no invalida la arquitectura; lo que demuestra es que la distribucion real tiene un costo de comunicacion que debe reconocerse en el analisis.",
    )
    add_para(
        doc,
        "Para trabajo futuro, conviene registrar diffs formales de salida, automatizar scripts de graficas, medir uso de memoria por fase, y probar recuperacion ante workers caidos. Tambien seria valioso comparar particionamientos alternativos manteniendo la propiedad fundamental: no separar puntos consecutivos del mismo bus y ruta.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
